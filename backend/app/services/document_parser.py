import io
import docx
import pdfplumber

def parse_docx(content: bytes) -> str:
    doc = docx.Document(io.BytesIO(content))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        table_text = _table_to_text(table)
        if table_text:
            paragraphs.append(table_text)

    return '\n\n'.join(paragraphs)

def parse_pdf(content: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                lines = page_text.split('\n')
                cleaned_lines = _clean_pdf_lines(lines)
                text_parts.append('\n'.join(cleaned_lines))

            for table in page.extract_tables():
                if table:
                    table_text = _table_to_text_str(table)
                    if table_text:
                        text_parts.append(table_text)

    return '\n\n'.join(text_parts)

def parse_txt(content: bytes) -> str:
    return content.decode('utf-8', errors='replace')

def _table_to_text(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    return _table_to_text_str(rows)

def _table_to_text_str(rows: list[list[str]]) -> str:
    if not rows or len(rows) < 2:
        return ""

    headers = rows[0]
    result_parts = []
    for row in rows[1:]:
        parts = []
        for i, cell in enumerate(row):
            if i < len(headers) and cell:
                parts.append(f"{headers[i]}{cell}")
        if parts:
            result_parts.append("，".join(parts))

    if result_parts:
        return "本表包含以下数据：" + "；".join(result_parts) + "。"
    return ""

def _clean_pdf_lines(lines: list[str]) -> list[str]:
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if len(stripped) < 5 and not any(c.isalpha() for c in stripped):
            continue
        cleaned.append(stripped)
    return cleaned
